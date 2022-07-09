"""
Author: Udayshankar R
GitHub: https://github.com/Uralstech

Helps WinSurfer with text and Python files!

"""

from FileAccessors.WSEFOpener import read, write
from subprocess import Popen, PIPE
from ShortCode.FileSystem import *
from ShortCode.InDev.UI import *
from tkinter import filedialog
from docx import Document
from USER import USER
import os

def open_file(root, current_path:str, data_path:str):
    if data_path in current_path:
        showerror("WinSurfer TextEditor", "Lol you aren't allowed to open files from here.")
        return
    else:
        test_path = ""
        for i in current_path:
            if i != "/":
                test_path += i
            else:
                test_path += "\\"
        current_path = test_path
        if data_path in current_path:
            showerror("WinSurfer TextEditor", "Lol you aren't allowed to open files from here.")
            return

    root2 = Toplevel(root)
    root2.title(os.path.basename(current_path))
    root2.geometry("650x400")
    root2.minsize(650, 400)

    file_type = os.path.splitext(current_path)[1]

    global root3
    root3 = None

    global main_font_size
    main_font_size = USER.logged_user.formatted_font[0]
    global main_font_family
    main_font_family = USER.logged_user.formatted_font[1]
    global main_font_colour
    main_font_colour = USER.logged_user.formatted_font[2]

    global console_font_size
    console_font_size = USER.logged_user.formatted_font[3]
    global console_font_family
    console_font_family = USER.logged_user.formatted_font[4]
    global console_font_colour
    console_font_colour = USER.logged_user.formatted_font[5]

    global main_font
    main_font = GetFont(main_font_family, main_font_size)
    global console_font
    console_font = GetFont(console_font_family, console_font_size)

    global main_tab
    main_tab = main_font.measure('    ')
    global console_tab
    console_tab = console_font.measure('    ')

    def change_font(main=True):
        global root3
        root3 = Toplevel(root2, bg="Black")
        root3.title("WinSurfer View Settings")
        root3.geometry("300x300")
        root3.resizable(False, False)

        def set_main_size(v):
            global main_font
            global main_font_size
            global main_tab
            main_font_size = s.get()

            main_font = font = GetFont(main_font_family, main_font_size)
            main_tab = tab = font.measure('    ')
            text['font'] = font
            text['tabs'] = tab
            l['font'] = font

            USER.logged_user.write_font([main_font_size, main_font_family, main_font_colour, console_font_size, console_font_family, console_font_colour])
            USER.update_user_data()
        def set_main_font():
            global main_font
            global main_font_family
            global main_font_colour
            global main_tab
            main_font_family = click2.get()
            main_font_colour = click1.get()

            main_font = font = GetFont(main_font_family, main_font_size)
            main_tab = tab = font.measure('    ')
            text['fg'] = main_font_colour
            text['font'] = font
            text['tabs'] = tab
            l['fg'] = main_font_colour
            l['font'] = font

            USER.logged_user.write_font([main_font_size, main_font_family, main_font_colour, console_font_size, console_font_family, console_font_colour])
            USER.update_user_data()

        def set_console_size(v):
            global console_font
            global console_font_size
            global console_tab
            console_font_size = s.get()

            console_font = font = GetFont(console_font_family, console_font_size)
            console_tab = tab = font.measure('    ')
            console['font'] = font
            console['tabs'] = tab
            l['font'] = font

            USER.logged_user.write_font([main_font_size, main_font_family, main_font_colour, console_font_size, console_font_family, console_font_colour])
            USER.update_user_data()
        def set_console_font():
            global console_font
            global console_font_family
            global console_font_colour
            global console_tab
            console_font_family = click2.get()
            console_font_colour = click1.get()
            console_font = font = GetFont(console_font_family, console_font_size)
            console_tab = tab = font.measure('    ')
            console['fg'] = console_font_colour
            console['font'] = font
            console['tabs'] = tab
            l['fg'] = console_font_colour
            l['font'] = font

            USER.logged_user.write_font([main_font_size, main_font_family, main_font_colour, console_font_size, console_font_family, console_font_colour])
            USER.update_user_data()

        tfont = GetFont("Terminal", 15)
        l = GetLabel(root3, "Text", bgcolour="black") ; l.pack(pady=10, side=TOP)
        s = Scale(root3, bg="Black", troughcolor="Black", font=tfont, fg="White", orient=HORIZONTAL, from_=10, to=30, command=set_main_size) ; s.pack(fill="x", expand=1, padx=10, anchor=S) ; s.set(main_font_size)
        d1, click1 = GetDropDown(root3, ["White", "Red", "Yellow", "Green", "Cyan", "Blue", "Purple"], 23, font=tfont, bgColour="Black", fgColour="White") ; d1.pack(pady=10, anchor=S) ; click1.set(main_font_colour)
        d2, click2 = GetDropDown(root3, ["Calibri", "Arial", "Arial Black", "Bahnschrift", "Bahnschrift Condensed", "Bodoni Bd BT", "Calibri", "Cambria", "Cascadia Code", "Comic Sans MS", "Consolas", "Courier", "Exotc350 Bd BT", "Impact", "Ink Free", "Lucida Console", "MS Gothic", "MS Sans Serif", "Microsoft Sans Serif", "Microsoft YaHei UI", "Nirmala UI", "OCR-A BT", "Segoe UI", "Segoe UI Black", "SimSun", "Square721 BT", "Terminal", "Times New Roman", "Verdana"], 23, font=tfont, bgColour="Black", fgColour="White") ; d2.pack(pady=10, anchor=S) ; click2.set(main_font_family)
        b = GetButton(root3, "Update font", width=27, font=tfont, bgColour="Black", fgColour="White", function=set_main_font)
        b.pack(pady=10, anchor=S)
        
        if main:
            click1.set(main_font_colour)
            click2.set(main_font_family)
            l['fg'] = main_font_colour
            l['font'] = main_font
        else:
            click1.set(console_font_colour)
            click2.set(console_font_family)
            l['fg'] = console_font_colour
            l['font'] = console_font
            s.set(console_font_size)

            s.config(command=set_console_size)
            b.config(command=set_console_font)
        
        root3.mainloop()

    def save_file(v=None, saveas=False):
        ftype = file_type
        cpath = current_path

        if saveas:
            types = (('Text file', '*.txt'),
                     ('Text Document (Buggy, not recommended)', '*.docx'),
                     ('WinSurfer Encrypted file', '*.wsef'),
                     ('Python file', '*.py'),
                     ('All files', '*.*'))
            
            base_path = r"C:\Users" + "\\" + os.getlogin()
            cpath = filedialog.asksaveasfilename(title="Create file", initialdir=base_path, confirmoverwrite=False, filetypes=types, defaultextension=types)

            edpath = ""
            for i in cpath:
                if i != "/": edpath += i
                else: edpath += "\\"
            cpath = edpath
            ftype = os.path.splitext(cpath)[1]

            if not cpath.startswith("C:\\") or file_exists(cpath):
                showerror("WinSurfer TextEditor", f"There was an error while creating the file at {cpath}. The path is empty or the file already exists.")
                return

        if ftype == ".wsef":
            write(cpath, list(text.get('0.0', 'end').split("\n")))
        elif ftype == ".docx":
            doc = Document(cpath)
            for i in doc.paragraphs:
                p = i._element
                p.getparent().remove(p)
                p._p = p._element = None

            for i in list(text.get('0.0', 'end').split("\n")):
                doc.add_paragraph(i)
            doc.save(cpath)
        else:
            file_override(cpath, text.get('0.0', 'end'))
        
        if saveas:
            open_file(root, cpath, data_path)

    def clear_console(v=None):
        console['state'] = 'normal'
        console.delete('0.0', 'end')
        console['state'] = 'disabled'
    def run_python(v=None):
        proccess = Popen(['python', current_path], stdout=PIPE, stderr=PIPE, shell=True)
        clear_console()

        out, err = proccess.communicate()
        console['state'] = 'normal'
        console.insert('0.0', out)
        console.insert('1.0', err)
        console['state'] = 'disabled'

    def close_window():
        if root3 != None:
            root3.destroy()
        root2.destroy()

    console = None
    if file_type == ".py":
        root2.bind('<F5>', run_python)
        root2.bind('<Control-c>', clear_console)
        GetMenu(root2, subcommands={"File" : {"Save" : save_file, "Save as" : lambda:save_file(saveas=True)}, "View" : {"Main" : change_font, "Console" : lambda:change_font(False)}, "Python" : {"Clear console" : clear_console, "Run file" : run_python}})
        console = Text(root2, tabs=console_tab, wrap=WORD, bg="Black", height=5, highlightthickness=2, highlightcolor="White", fg=console_font_colour, font=console_font)
        console.insert('0.0', 'Thanks for using WinSurfer as your IDE!\nThis will be your console.\nNOTE: os.getcwd() will return the folder that WinSurfer is in.\nNOTE: input() won\'t work.')
        console.pack(fill="both", expand=1, side=TOP)
        console['state'] = 'disabled'
    else:
        GetMenu(root2, subcommands={"File" : {"Save" : save_file, "Save as" : lambda:save_file(saveas=True)}, "View" : {"Main" : change_font}})

    text = Text(root2, tabs=main_tab, wrap=NONE, bg="Black", fg=main_font_colour, font=main_font, insertbackground="White")
    text.pack(fill="both", expand=1, side=TOP)
    if file_type == ".wsef":
        file = read(current_path, format_string=True)
        fstr = ""
        for i in file:
            fstr += i
        text.insert('0.0', fstr)
    elif file_type == ".docx":
        doc = Document(current_path)
        fstr = ""
        for i in doc.paragraphs:
            fstr += i.text + "\n"
        text.insert('0.0', fstr)
    else:
        file = file_readlines(current_path)
        fstr = ""
        for i in file:
            fstr += i
        text.insert('0.0', fstr)

    root2.bind('<Control-s>', save_file)
    root2.protocol("WM_DELETE_WINDOW", close_window)
    root2.mainloop()
